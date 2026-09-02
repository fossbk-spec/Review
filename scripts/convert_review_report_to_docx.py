#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
convert_review_report_to_docx.py — Chuyển đổi báo cáo phản biện FINAL_REVIEW_REPORT.md
sang định dạng Word (.docx) chuẩn format quốc tế của Springer Nature.
"""
import sys
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")


def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def build_docx_report(md_path: Path, docx_path: Path):
    doc = docx.Document()

    # Page setup — 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header.is_linked_to_previous = False
        p_hdr = section.header.paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_hdr = p_hdr.add_run("Discover Artificial Intelligence | Peer Review Report")
        r_hdr.font.name = "Calibri"
        r_hdr.font.size = Pt(8.5)
        r_hdr.font.color.rgb = RGBColor(128, 128, 128)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("PEER REVIEW REPORT")
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0, 51, 102)  # Springer Navy

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("SPRINGER NATURE DISCOVER JOURNALS SERIES")
    r_sub.font.name = "Calibri"
    r_sub.font.size = Pt(11)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(80, 80, 80)

    # Metadata Table
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_widths = [Inches(1.8), Inches(4.7)]
    meta_data = [
        ("Journal", "Discover Artificial Intelligence"),
        ("Manuscript ID", "1cac0119-f585-45d1-87aa-65c7fecddf40"),
        ("Article Title", "Learning Analytics for Detecting Digital Information Overload Among Postgraduate Students Using Machine Learning"),
        ("Submission Round", "Revised Submission (Post Minor Revision)"),
        ("Review Guidelines", "https://link.springer.com/brands/discover/for-reviewers (Sound Science / Technical Validity)"),
    ]

    for i, (label, val) in enumerate(meta_data):
        row = table.rows[i]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width, c1.width = col_widths[0], col_widths[1]
        set_cell_background(c0, "F0F4F8")
        set_cell_background(c1, "FAFAFA")
        set_cell_margins(c0, top=100, bottom=100, left=140, right=140)
        set_cell_margins(c1, top=100, bottom=100, left=140, right=140)

        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        r0.font.name = "Calibri"
        r0.font.size = Pt(9.5)
        r0.font.bold = True
        r0.font.color.rgb = RGBColor(0, 51, 102)

        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(val)
        r1.font.name = "Calibri"
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGBColor(40, 40, 40)
        if label == "Article Title":
            r1.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Content Parsing
    raw_lines = md_path.read_text(encoding="utf-8").splitlines()
    in_part = False

    for line in raw_lines:
        line_s = line.strip()
        if not line_s:
            continue

        # Skip header metadata lines already in the table
        if line_s.startswith("# PEER REVIEW REPORT") or line_s.startswith("> **") or line_s == "---":
            continue

        # Level 1: PART 1 / PART 2
        if line_s.startswith("## PART 1:") or line_s.startswith("## PART 2:"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(line_s.replace("## ", ""))
            r.font.name = "Calibri"
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0, 51, 102)
            continue

        # Level 2 Headings: ### 1. Key Results...
        if line_s.startswith("### "):
            heading_text = line_s.replace("### ", "")
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(heading_text)
            r.font.name = "Calibri"
            r.font.size = Pt(11.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(30, 30, 30)

            # Special box for Recommendation under PART 2
            if "Recommendation" in heading_text:
                pass
            continue

        # Bullet lists
        if line_s.startswith("* ") or line_s.startswith("- "):
            bullet_text = line_s[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15

            # Process inline bolding: * **Label:** Text
            parts = bullet_text.split("**")
            if len(parts) >= 3:
                # Part 0 is prefix, Part 1 is bold, Part 2 is rest
                if parts[0]:
                    r = p.add_run(parts[0])
                    r.font.name = "Calibri"
                    r.font.size = Pt(10.5)
                r_bold = p.add_run(parts[1])
                r_bold.font.name = "Calibri"
                r_bold.font.size = Pt(10.5)
                r_bold.font.bold = True
                r_bold.font.color.rgb = RGBColor(0, 51, 102)

                r_rest = p.add_run("".join(parts[2:]))
                r_rest.font.name = "Calibri"
                r_rest.font.size = Pt(10.5)
                r_rest.font.color.rgb = RGBColor(50, 50, 50)
            else:
                r = p.add_run(bullet_text)
                r.font.name = "Calibri"
                r.font.size = Pt(10.5)
                r.font.color.rgb = RGBColor(50, 50, 50)
            continue

        # Numbered list: 1. ... 2. ...
        if line_s[0].isdigit() and line_s[1:3] in (". ", "):"):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            num_text = line_s[3:].strip()

            parts = num_text.split("**")
            if len(parts) >= 3:
                r_bold = p.add_run(parts[1])
                r_bold.font.name = "Calibri"
                r_bold.font.size = Pt(10.5)
                r_bold.font.bold = True
                r_bold.font.color.rgb = RGBColor(0, 51, 102)

                r_rest = p.add_run("".join(parts[2:]))
                r_rest.font.name = "Calibri"
                r_rest.font.size = Pt(10.5)
                r_rest.font.color.rgb = RGBColor(50, 50, 50)
            else:
                r = p.add_run(num_text)
                r.font.name = "Calibri"
                r.font.size = Pt(10.5)
                r.font.color.rgb = RGBColor(50, 50, 50)
            continue

        # Recommendation Highlight Box
        if line_s.startswith("**Minor Revision"):
            box_table = doc.add_table(rows=1, cols=1)
            box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = box_table.rows[0].cells[0]
            cell.width = Inches(6.5)
            set_cell_background(cell, "EBF3FA")
            set_cell_margins(cell, top=120, bottom=120, left=180, right=180)

            bp = cell.paragraphs[0]
            bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp.paragraph_format.space_after = Pt(0)
            br = bp.add_run("RECOMMENDED DECISION: MINOR REVISION (Essential Clarifications)")
            br.font.name = "Calibri"
            br.font.size = Pt(11)
            br.font.bold = True
            br.font.color.rgb = RGBColor(0, 51, 102)
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        # Checkbox lines: [x] or [ ]
        if line_s.startswith("[x]") or line_s.startswith("[ ]") or line_s.startswith("- [x]") or line_s.startswith("- [ ]"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            checked = "[x]" in line_s
            text_part = line_s.replace("- [x]", "").replace("- [ ]", "").replace("[x]", "").replace("[ ]", "").strip()

            r_box = p.add_run("☑  " if checked else "☐  ")
            r_box.font.name = "Segoe UI Symbol"
            r_box.font.size = Pt(11)
            r_box.font.bold = True
            r_box.font.color.rgb = RGBColor(0, 102, 204) if checked else RGBColor(128, 128, 128)

            parts = text_part.split("**")
            if len(parts) >= 3:
                r_b = p.add_run(parts[1])
                r_b.font.name = "Calibri"
                r_b.font.size = Pt(10)
                r_b.font.bold = True
                r_r = p.add_run("".join(parts[2:]))
                r_r.font.name = "Calibri"
                r_r.font.size = Pt(10)
            else:
                r_txt = p.add_run(text_part)
                r_txt.font.name = "Calibri"
                r_txt.font.size = Pt(10)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15

        # Format italics/bold
        r = p.add_run(line_s)
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(50, 50, 50)

    # Save document
    doc.save(str(docx_path))
    print(f"[+] Da tao thanh cong file docx tai: {docx_path}")


if __name__ == "__main__":
    md_file = Path("C:/Antigravity/Academic_Peer_Review/reviews/2026/REV_discover_ai_1cac0119/FINAL_REVIEW_REPORT.md")
    docx_file = Path("C:/Antigravity/Academic_Peer_Review/reviews/2026/REV_discover_ai_1cac0119/FINAL_REVIEW_REPORT.docx")

    if len(sys.argv) > 2:
        md_file = Path(sys.argv[1])
        docx_file = Path(sys.argv[2])

    build_docx_report(md_file, docx_file)
