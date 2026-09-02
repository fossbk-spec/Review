#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
anonymize_manuscript.py — Cong cu HO TRO an danh hoa ban thao truoc khi trich
xuat noi dung (KHONG thay the viec tu kiem tra bang mat).

Xu ly: doc file .docx/.pdf/.txt/.md, chuyen sang text/markdown:
- UU TIEN 1: doc file tu manuscripts/clean_markdown/ (neu da qua batch_ingestion.py)
- UU TIEN 2: file .md/.txt co san trong manuscripts/raw/ (co canh bao chua chuan hoa)
- UU TIEN 3: tu dong Ingest native bang python-docx / pypdf (chuan hdbs-anydoc, 0 token AI)
- UU TIEN 4 (du phong): dung pandoc ngoai neu co san.

Sau do xoa/che cac mau hinh lo danh tinh (PII):
- Email, ORCID, Phone/Fax
- Dong "Author(s):", "Tác giả:", "Affiliation:", "Corresponding author", "Grant/Funding"
- Ten file goc (thuong chua ten tac gia).

GIOI HAN QUAN TRONG (doc truoc khi dung that):
- Chi bat duoc mau hinh PHO BIEN bang regex - khong bat duoc moi truong hop
  (vi du ten tac gia xuat hien tu nhien trong cau van, watermark trong PDF,
  metadata an trong file .docx/.pdf nhu "Author" property).
- Anh/hinh ve co the chua watermark/logo vien nghien cuu - script nay KHONG xu ly anh.
- BAT BUOC tu doc lai file da an danh bang mat truoc khi tin tuong no da sach.
"""
import argparse
import re
import shutil
import subprocess
import sys
from pathlib import Path

# Chong loi font chu Unicode tren Windows Terminal (cp1252)
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

EMAIL_RE = re.compile(r'\b[\w.+-]+@[\w-]+\.[\w.-]+\b')
ORCID_RE = re.compile(r'https?://orcid\.org/\d{4}-\d{4}-\d{4}-[\dX]{4}', re.IGNORECASE)

IDENTIFYING_LINE_PATTERNS = [
    re.compile(r'(?im)^.*\bauthor\s*\(?s?\)?\s*:.*$'),
    re.compile(r'(?im)^.*\btác giả\s*:.*$'),
    re.compile(r'(?im)^.*\b(affiliation|đơn vị|viện nghiên cứu)s?\s*:.*$'),
    re.compile(r'(?im)^.*\bcorresponding author\b.*$'),
    re.compile(r'(?im)^.*\b(grant|funding)\s*(no\.?|number)?\s*:.*$'),
    re.compile(r'(?im)^.*\b(tel|phone|fax)\s*:.*$'),
]

METADATA_STRIP_FIELDS = ["author", "title", "creator", "producer", "lastmodifiedby"]
MOJIBAKE_SIGNATURES = ["\ufffd", "â€™", "â€œ", "â€", "Ã¡", "Ã©", "Ã¢", "Ã³", "Ãº", "Ä‘"]


def _extract_docx_native(file_path: Path) -> str:
    """Boc tach DOCX sang Markdown bang python-docx (dong co hdbs-anydoc)."""
    try:
        import docx
    except ImportError:
        return ""

    doc = docx.Document(str(file_path))
    lines = []
    for element in doc.element.body:
        tag = element.tag.split('}')[-1]
        if tag == 'p':
            for p in doc.paragraphs:
                if p._p == element:
                    text = p.text.strip()
                    if not text:
                        continue
                    style = p.style.name.lower()
                    if "heading 1" in style or style == "title":
                        lines.append(f"\n# {text}\n")
                    elif "heading 2" in style or style == "subtitle":
                        lines.append(f"\n## {text}\n")
                    elif "heading 3" in style:
                        lines.append(f"\n### {text}\n")
                    elif "heading 4" in style:
                        lines.append(f"\n#### {text}\n")
                    elif "list" in style or style.startswith("bullet"):
                        lines.append(f"* {text}")
                    else:
                        lines.append(f"\n{text}\n")
                    break
        elif tag == 'tbl':
            for table in doc.tables:
                if table._tbl == element:
                    table_rows = []
                    for row in table.rows:
                        row_cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                        table_rows.append(row_cells)
                    if table_rows:
                        header = table_rows[0]
                        lines.append("\n| " + " | ".join(header) + " |")
                        lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                        for r in table_rows[1:]:
                            padded_r = r + [""] * (len(header) - len(r))
                            lines.append("| " + " | ".join(padded_r[:len(header)]) + " |")
                        lines.append("\n")
                    break
    return "\n".join(lines).strip()


def _extract_pdf_native(file_path: Path) -> tuple[str, bool]:
    """Boc tach PDF sang Markdown bang pypdf kem kiem tra mojibake."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return "", False

    reader = PdfReader(str(file_path))
    total_pages = len(reader.pages)
    lines = [f"# Bản thảo trích xuất từ: {file_path.name}", f"> Tổng số trang: {total_pages}\n"]
    has_mojibake = False

    for idx, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text:
            clean_text = text.replace("\r\n", "\n").strip()
            for sig in MOJIBAKE_SIGNATURES:
                if sig in clean_text:
                    has_mojibake = True
                    break
            lines.append(f"\n<!-- Trang {idx}/{total_pages} -->")
            lines.append(f"### [Trang {idx}]\n")
            lines.append(clean_text)
            lines.append("\n---\n")

    return "\n".join(lines).strip(), has_mojibake


def convert_to_markdown(src: Path, tmp_dir: Path) -> tuple[str, str]:
    """Chuyen doi sang markdown. Tra ve (text, nguon).
    nguon in {"anydoc_clean", "native_text", "native_ingest", "pandoc_fallback"} — 4 trang thai
    RIENG BIET, minh bach ve nguon goc du lieu:

    UU TIEN 1: co ban anydoc/hdbs-anydoc san sinh o manuscripts/clean_markdown/ -> "anydoc_clean"
    UU TIEN 2: file da la .md/.txt co san trong raw/ -> "native_text"
    UU TIEN 3: tu dong parse native bang python-docx / pypdf -> "native_ingest"
    UU TIEN 4 (du phong): dung pandoc truc tiep -> "pandoc_fallback"
    """
    # UU TIEN 1: Kiem tra ban sach tu batch_ingestion
    clean_md_path = Path("manuscripts/clean_markdown") / (src.stem + ".md")
    if clean_md_path.exists():
        return clean_md_path.read_text(encoding="utf-8", errors="replace"), "anydoc_clean"

    # UU TIEN 2: File van ban co san
    ext = src.suffix.lower()
    if ext in (".md", ".txt"):
        print(f"  [!] Luu y: {src.name} la file text/markdown co san, khong tim thay ban")
        print(f"      tuong ung da qua anydoc o {clean_md_path}. Neu day la ban thao tho")
        print(f"      chua qua ingestion, van co the con rac dinh dang chua duoc chuan hoa.")
        return src.read_text(encoding="utf-8", errors="replace"), "native_text"

    # UU TIEN 3: Tu dong Ingest native bang dong co hdbs-anydoc (docx / pypdf)
    if ext == ".docx":
        docx_text = _extract_docx_native(src)
        if docx_text:
            return docx_text, "native_ingest"

    if ext == ".pdf":
        pdf_text, has_mojibake = _extract_pdf_native(src)
        if pdf_text:
            if has_mojibake:
                warning_header = (
                    "<!-- CẢNH BÁO HDBS-ANYDOC: Phát hiện dấu hiệu lỗi font/mojibake tiếng Việt trong PDF gốc. -->\n"
                    "<!-- Khuyến nghị: Ưu tiên dùng bản .docx hoặc đối chiếu thủ công với bản gốc. -->\n\n"
                )
                pdf_text = warning_header + pdf_text
            return pdf_text, "native_ingest"

    # UU TIEN 4: Du phong cuoi cung voi Pandoc (neu co)
    print(f"  [!] CANH BAO: chua tim thay {clean_md_path}")
    print(f"      Dang thu dung pandoc lam DU PHONG ngoai...")

    if shutil.which("pandoc"):
        out = tmp_dir / (src.stem + ".md")
        result = subprocess.run(
            ["pandoc", str(src), "-t", "markdown", "-o", str(out)],
            capture_output=True, text=True)
        if result.returncode == 0 and out.exists():
            return out.read_text(encoding="utf-8", errors="replace"), "pandoc_fallback"
        print(f"  [!] pandoc loi ({result.stderr.strip()[:200]}) - thu cach khac.")

    raise RuntimeError(
        f"Khong the chuyen doi {src.name} - hay cai dat python-docx va pypdf "
        f"('pip install python-docx pypdf') hoac chay anydoc/hdbs-anydoc truoc."
    )


def anonymize_text(text: str) -> tuple[str, dict]:
    stats = {
        "emails_removed": 0, 
        "orcids_removed": 0, 
        "identifying_lines_removed": 0
    }

    def _count_and_replace(pattern, replacement, s, key):
        matches = pattern.findall(s)
        stats[key] += len(matches)
        return pattern.sub(replacement, s)

    text = _count_and_replace(EMAIL_RE, "[EMAIL ĐÃ ẨN]", text, "emails_removed")
    text = _count_and_replace(ORCID_RE, "[ORCID ĐÃ ẨN]", text, "orcids_removed")
    for pat in IDENTIFYING_LINE_PATTERNS:
        text = _count_and_replace(pat, "[DÒNG NHẬN DIỆN ĐÃ ẨN]", text, "identifying_lines_removed")

    return text, stats


def main() -> int:
    parser = argparse.ArgumentParser(description="Ho tro an danh hoa ban thao truoc khi dua vao AI.")
    parser.add_argument("input", type=Path, help="File trong manuscripts/raw/")
    parser.add_argument("-o", "--output", type=Path, default=None,
                         help="Mac dinh: manuscripts/anonymized/<ten>.md")
    args = parser.parse_args()

    if not args.input.exists():
        print(f"[-] Khong tim thay file: {args.input}")
        return 1

    output = args.output or Path("manuscripts/anonymized") / (args.input.stem + ".md")
    output.parent.mkdir(parents=True, exist_ok=True)

    tmp_dir = Path("manuscripts/anonymized/.tmp")
    tmp_dir.mkdir(parents=True, exist_ok=True)

    try:
        raw_text, source = convert_to_markdown(args.input, tmp_dir)
    except RuntimeError as e:
        print(f"[-] {e}")
        return 1
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    clean_text, stats = anonymize_text(raw_text)
    output.write_text(clean_text, encoding="utf-8")

    source_label = {
        "anydoc_clean": "anydoc/hdbs-anydoc (đúng khuyến nghị - từ manuscripts/clean_markdown/)",
        "native_text": "file text/markdown gốc — CHƯA XÁC NHẬN đã qua anydoc",
        "native_ingest": "hdbs-anydoc native (docx/pypdf tự động - không cần Pandoc ngoài)",
        "pandoc_fallback": "pandoc dự phòng — CHƯA qua anydoc, xem cảnh báo phía trên",
    }[source]

    print(f"\n========================================================")
    print(f" [ẨN DANH HÓA BẢN THẢO] HOÀN TẤT")
    print(f"========================================================")
    print(f"[+] Da ghi:  {output}")
    print(f"[+] Nguon:   {source_label}")
    print(f"[+] Da an:   {stats['emails_removed']} email, {stats['orcids_removed']} ORCID, "
          f"{stats['identifying_lines_removed']} dong nhan dien (Author/Affiliation/Tel...).")
    print(f"========================================================")
    print(f"[!] BAT BUOC: mo lai file va tu doc kiem tra bang mat truoc khi tin da sach.")
    print(f"[!] Script nay KHONG xu ly: ten tac gia xuat hien tu nhien trong cau van,")
    print(f"    watermark/logo trong anh, metadata an trong file goc (Author property).")
    print(f"    Neu file goc la PDF/DOCX, kiem tra rieng thuoc tinh 'Author' cua file")
    print(f"    (vi du: exiftool, hoac File > Properties trong Word/Adobe).\n")
    return 0


if __name__ == "__main__":
    sys.exit(main())
